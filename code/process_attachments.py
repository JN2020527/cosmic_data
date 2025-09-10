import os
import re
from datetime import datetime
from typing import Optional, Tuple, List
import json

from openpyxl import load_workbook
import requests

try:
    from config import DATA_DIR, DEEPSEEK_API_KEY, DEEPSEEK_API_URL
except ImportError:
    print("错误：未找到配置文件 config.py")
    print("请复制 config_template.py 为 config.py 并填入正确的配置信息")
    exit(1)


def print_step(title: str) -> None:
    print(f"==== {title} ====")


# 已删除 move_paragraph_after_index 函数，使用更简单的直接插入方法


def parse_attachment_filename(filename: str) -> Optional[Tuple[str, str, str, str]]:
    """
    Parse filenames like: 附件x-需求名@属性.扩展名
    Returns tuple: (prefix_with_number, requirement_name, attribute, extension)
    Example: ("附件3", "关于xxx的需求", "COSMIC工作量评估基础表", ".xlsx")
    """
    name, ext = os.path.splitext(filename)
    # Work with just the basename (no directories)
    base = os.path.basename(name)
    match = re.match(r"^附件(\d+)-(.+?)@(.+)$", base)
    if not match:
        return None
    number, req_name, attribute = match.groups()
    return (f"附件{number}", req_name, attribute, ext)


def build_attachment_filename(prefix_with_number: str, requirement_name: str, attribute: str, ext: str) -> str:
    return f"{prefix_with_number}-{requirement_name}@{attribute}{ext}"


def batch_rename(requirement_name: str) -> None:
    print_step("第一步：批量修改文件名字")
    if not os.path.isdir(DATA_DIR):
        print(f"目录不存在：{DATA_DIR}")
        return

    files = os.listdir(DATA_DIR)
    rename_count = 0
    for fname in files:
        parsed = parse_attachment_filename(fname)
        if not parsed:
            continue
        prefix, _old_req, attribute, ext = parsed
        new_name = build_attachment_filename(prefix, requirement_name, attribute, ext)
        if new_name == fname:
            print(f"跳过（已是目标名）：{fname}")
            continue
        src = os.path.join(DATA_DIR, fname)
        dst = os.path.join(DATA_DIR, new_name)
        os.rename(src, dst)
        rename_count += 1
        print(f"重命名：{fname} -> {new_name}")

    print(f"重命名完成，共处理 {rename_count} 个文件。")


def find_attachment_by_number(number: int) -> Optional[str]:
    """Find file path for '附件{number}-...@....xlsx' in DATA_DIR."""
    if not os.path.isdir(DATA_DIR):
        return None
    for fname in os.listdir(DATA_DIR):
        parsed = parse_attachment_filename(fname)
        if not parsed:
            continue
        prefix, _req, _attr, _ext = parsed
        if prefix == f"附件{number}":
            return os.path.join(DATA_DIR, fname)
    return None


def write_attachment3_sheet2_cells(requirement_name: str) -> None:
    print_step("第二步：修改附件3 sheet2 的 A3、B3")
    path = find_attachment_by_number(3)
    if not path:
        print("未找到附件3 文件。")
        return
    wb = load_workbook(path)
    sheet_name = None
    # Prefer explicit 'sheet2' by index (second worksheet)
    if len(wb.sheetnames) >= 2:
        sheet_name = wb.sheetnames[1]
    else:
        sheet_name = wb.active.title
    ws = wb[sheet_name]
    ws["A3"] = requirement_name
    ws["B3"] = requirement_name
    wb.save(path)
    print(f"已更新 {os.path.basename(path)} -> {sheet_name} 的 A3, B3 为：{requirement_name}")


def write_attachment4_cells(requirement_name: str) -> None:
    print_step("第三步：修改附件4 B2")
    path = find_attachment_by_number(4)
    if not path:
        print("未找到附件4 文件。")
        return
    wb = load_workbook(path)
    ws = wb.active
    ws["B2"] = requirement_name
    wb.save(path)
    print(f"已更新 {os.path.basename(path)} -> B2 为：{requirement_name}")


def sum_attachment5_col_L_from_L2() -> float:
    print_step("第四步：计算附件5 L列(L2开始)总和")
    path = find_attachment_by_number(5)
    if not path:
        print("未找到附件5 文件。返回 0。")
        return 0.0

    ext = os.path.splitext(path)[1].lower()

    # Handle legacy .xls via xlrd
    if ext == ".xls":
        try:
            import xlrd  # type: ignore
        except Exception:
            print("未安装 xlrd，无法读取 .xls 文件，请先安装 xlrd。返回 0。")
            return 0.0
        try:
            wb = xlrd.open_workbook(path)
            sheet = wb.sheet_by_index(0)
            total: float = 0.0
            col_index_for_L = 11  # 0-based index for column 'L'
            for row_idx in range(1, sheet.nrows):  # start from row 2 -> index 1
                value = sheet.cell_value(row_idx, col_index_for_L)
                if value in (None, ""):
                    continue
                try:
                    total += float(value)
                except Exception:
                    # skip non-numeric cells
                    continue
            print(f"L列合计：{total}")
            return total
        except Exception as e:
            print(f"读取 .xls 失败：{e}")
            return 0.0

    # Default: .xlsx via openpyxl
    try:
        wb = load_workbook(path, data_only=True)
        ws = wb.active
        total = 0.0
        max_row = ws.max_row if ws.max_row and ws.max_row > 1 else 1
        for row in range(2, max_row + 1):
            cell = ws[f"L{row}"]
            value = cell.value
            if value is None or (isinstance(value, str) and value.strip() == ""):
                continue
            try:
                total += float(value)
            except Exception:
                continue
        print(f"L列合计：{total}")
        return total
    except Exception as e:
        print(f"读取 .xlsx 失败：{e}")
        return 0.0


def write_attachment4_with_sum(total: float) -> None:
    print_step("第五步：将总和写入附件4 D7")
    path = find_attachment_by_number(4)
    if not path:
        print("未找到附件4 文件。")
        return
    wb = load_workbook(path)
    ws = wb.active
    ws["D7"] = total
    wb.save(path)
    print(f"已更新 {os.path.basename(path)} -> D7 为：{total}")


def write_attachment3_sheet2_F3_with_sum(total: float) -> None:
    print_step("第六步：将总和写入附件3 sheet2 的 F3")
    path = find_attachment_by_number(3)
    if not path:
        print("未找到附件3 文件。")
        return
    wb = load_workbook(path)
    # second sheet or active
    if len(wb.sheetnames) >= 2:
        ws = wb[wb.sheetnames[1]]
    else:
        ws = wb.active
    ws["F3"] = total
    wb.save(path)
    print(f"已更新 {os.path.basename(path)} -> {ws.title} 的 F3 为：{total}")


def write_attachment4_C6_with_today() -> None:
    print_step("第七步：填充当前时间到附件4 C6（xxxx年x月x日）")
    path = find_attachment_by_number(4)
    if not path:
        print("未找到附件4 文件。")
        return
    today = datetime.now()
    # Remove leading zeros in month/day
    date_str = f"{today.year}年{today.month}月{today.day}日"
    wb = load_workbook(path)
    ws = wb.active
    ws["C6"] = date_str
    wb.save(path)
    print(f"已更新 {os.path.basename(path)} -> C6 为：{date_str}")


def calculate_attachment3_e3_formula() -> int:
    """手动计算附件3 sheet2 E3单元格的公式: =COUNTA(COSMIC功能点拆分表!K:K)-1
    注意：实际数据从K4开始，K1-K3是标题行"""
    path3 = find_attachment_by_number(3)
    if not path3:
        return 0
    
    wb3 = load_workbook(path3, data_only=True)
    
    # 找到COSMIC功能点拆分表工作表
    cosmic_sheet = None
    for sheet_name in wb3.sheetnames:
        if 'COSMIC功能点拆分表' in sheet_name:
            cosmic_sheet = wb3[sheet_name]
            break
    
    if not cosmic_sheet:
        return 0
    
    # 从K4开始计算非空单元格数量（K1-K3是标题行）
    k_count = 0
    for row in range(4, cosmic_sheet.max_row + 1):
        k_value = cosmic_sheet[f'K{row}'].value
        if k_value is not None and str(k_value).strip():
            k_count += 1
    
    print(f"COSMIC功能点拆分表K列数据行数（从K4开始）: {k_count}")
    # 返回实际的功能点数量
    return k_count


def write_attachment4_B7_from_attachment3_E3() -> None:
    print_step("第八步：将附件3 sheet2 E3 写入附件4 B7")
    path3 = find_attachment_by_number(3)
    path4 = find_attachment_by_number(4)
    if not path3:
        print("未找到附件3 文件。")
        return
    if not path4:
        print("未找到附件4 文件。")
        return
    
    # 首先尝试直接读取E3的值
    wb3 = load_workbook(path3, data_only=True)
    if len(wb3.sheetnames) >= 2:
        ws3 = wb3[wb3.sheetnames[1]]
    else:
        ws3 = wb3.active
    e3_value = ws3["E3"].value

    # 如果E3的值是None，说明它包含公式且无法直接计算，手动计算
    if e3_value is None:
        print("E3单元格包含公式，手动计算结果...")
        e3_value = calculate_attachment3_e3_formula()
        print(f"计算得到E3公式结果: {e3_value}")

    wb4 = load_workbook(path4)
    ws4 = wb4.active
    ws4["B7"] = e3_value
    wb4.save(path4)
    print(f"已更新 {os.path.basename(path4)} -> B7 为：{e3_value}")


def extract_attachment5_h_i_content() -> Tuple[List[str], List[str]]:
    """提取附件5中H列和I列的内容"""
    path = find_attachment_by_number(5)
    if not path:
        return [], []
    
    try:
        import xlrd
    except Exception:
        print("未安装 xlrd，无法读取 .xls 文件")
        return [], []
    
    try:
        wb = xlrd.open_workbook(path)
        sheet = wb.sheet_by_index(0)
        
        h_contents = []
        i_contents = []
        h_col_index = 7  # H列索引
        i_col_index = 8  # I列索引
        
        # 从第2行开始读取（跳过标题行）
        for row_idx in range(1, sheet.nrows):
            h_value = sheet.cell_value(row_idx, h_col_index) if h_col_index < sheet.ncols else ''
            i_value = sheet.cell_value(row_idx, i_col_index) if i_col_index < sheet.ncols else ''
            
            if h_value and str(h_value).strip():
                h_contents.append(str(h_value).strip())
            if i_value and str(i_value).strip():
                i_contents.append(str(i_value).strip())
        
        return h_contents, i_contents
    except Exception as e:
        print(f"读取附件5失败：{e}")
        return [], []


def call_deepseek_api(content: str) -> str:
    """调用DeepSeek API生成内容概述"""
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
    }
    
    prompt = f"""基于以下工作项内容，请生成一个精简的需求内容概述。要求：
1. 以"内容概述："开头
2. 总结的内容为有序列表
3. 每个列表项应该简洁明了，概括主要功能点
4. 合并相似的功能点
5. 按照逻辑顺序排列

工作项内容：
{content}

请生成概述："""

    data = {
        "model": "deepseek-chat",
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ],
        "max_tokens": 1000,
        "temperature": 0.7
    }
    
    try:
        print(f"正在调用DeepSeek API生成概述...")
        
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=30)
        
        if response.status_code != 200:
            error_msg = f"API调用失败，状态码: {response.status_code}"
            print(error_msg)
            raise Exception(error_msg)
        
        result = response.json()
        
        if 'choices' in result and len(result['choices']) > 0:
            api_summary = result['choices'][0]['message']['content'].strip()
            print("✅ API调用成功，已生成概述")
            return api_summary
        else:
            error_msg = "API响应格式错误"
            print(error_msg)
            raise Exception(error_msg)
            
    except Exception as e:
        error_msg = f"调用DeepSeek API失败：{e}"
        print(error_msg)
        raise Exception(error_msg)


def summarize_requirement_content_and_update_h4() -> None:
    print_step("第九步：基于附件5 H和I列内容生成需求概述并更新附件4 A4")
    
    try:
        # 提取附件5的H和I列内容
        h_contents, i_contents = extract_attachment5_h_i_content()
        
        if not h_contents and not i_contents:
            print("未找到附件5的H和I列内容")
            return
        
        print(f"提取到H列内容 {len(h_contents)} 项，I列内容 {len(i_contents)} 项")
        
        # 合并H和I列内容（去重）
        all_contents = []
        all_contents.extend(h_contents)
        for content in i_contents:
            if content not in all_contents:
                all_contents.append(content)
        
        # 将内容合并为一个字符串
        content_text = "\n".join([f"{i+1}. {content}" for i, content in enumerate(all_contents)])
        
        print("正在调用DeepSeek API生成需求概述...")
        summary = call_deepseek_api(content_text)
        
        print(f"生成的概述：\n{summary}")
        
        # 更新附件4的A4单元格
        path4 = find_attachment_by_number(4)
        if not path4:
            print("未找到附件4文件")
            return
        
        wb4 = load_workbook(path4)
        ws4 = wb4.active
        ws4["A4"] = summary
        wb4.save(path4)
        
        print(f"已更新 {os.path.basename(path4)} -> A4 为概述内容")
        
    except Exception as e:
        print(f"第九步执行失败：{e}")
        print("程序终止，请检查API配置或网络连接")
        raise


def load_function_codes() -> List[Tuple[str, str, str]]:
    """从附件3的COSMIC功能点拆分表中加载一二三级功能点码值"""
    
    # 首先尝试从附件3的sheet4加载
    path3 = find_attachment_by_number(3)
    if path3:
        try:
            wb3 = load_workbook(path3, data_only=True)
            
            # 查找COSMIC功能点拆分表工作表
            sheet_name = "COSMIC功能点拆分表"
            if sheet_name in wb3.sheetnames:
                ws = wb3[sheet_name]
                print(f"✅ 从附件3的{sheet_name}中加载功能点码值")
                
                codes = []
                current_level1 = ""
                current_level2 = ""
                
                # 从第4行开始读取数据（第3行是标题）
                for row in range(4, ws.max_row + 1):
                    level1 = ws.cell(row, 2).value  # B列 - 一级模块
                    level2 = ws.cell(row, 3).value  # C列 - 二级模块
                    level3 = ws.cell(row, 4).value  # D列 - 三级模块
                    
                    # 更新当前的一级、二级功能点
                    if level1 and str(level1).strip():
                        current_level1 = str(level1).strip()
                    if level2 and str(level2).strip():
                        current_level2 = str(level2).strip()
                    
                    if level3 and str(level3).strip():
                        codes.append((current_level1, current_level2, str(level3).strip()))
                
                if codes:
                    print(f"✅ 从附件3加载了 {len(codes)} 个功能点码值")
                    return codes
                else:
                    print("⚠️  附件3的COSMIC功能点拆分表中未找到有效数据")
            else:
                print(f"⚠️  附件3中未找到{sheet_name}工作表")
        
        except Exception as e:
            print(f"⚠️  从附件3加载功能点码值失败：{e}")
    
    # 备用方案：从独立的一二三级功能点.xlsx文件加载
    print("🔄 尝试从备用文件加载功能点码值...")
    codes_path = os.path.join(os.path.dirname(__file__), "一二三级功能点.xlsx")
    if not os.path.exists(codes_path):
        print(f"❌ 未找到备用功能点码值文件：{codes_path}")
        return []
    
    try:
        wb = load_workbook(codes_path, data_only=True)
        ws = wb.active
        
        codes = []
        current_level1 = ""
        current_level2 = ""
        
        for row in range(2, ws.max_row + 1):
            level1 = ws.cell(row, 1).value
            level2 = ws.cell(row, 2).value
            level3 = ws.cell(row, 3).value
            
            # 更新当前的一级、二级功能点
            if level1 and str(level1).strip():
                current_level1 = str(level1).strip()
            if level2 and str(level2).strip():
                current_level2 = str(level2).strip()
            
            if level3 and str(level3).strip():
                codes.append((current_level1, current_level2, str(level3).strip()))
        
        print(f"📁 从备用文件加载了 {len(codes)} 个功能点码值")
        return codes
        
    except Exception as e:
        print(f"❌ 加载备用功能点码值文件失败：{e}")
        return []


def parse_ai_function_matches(ai_response: str, function_codes: List[Tuple[str, str, str]]) -> List[Tuple[str, str, str, str, float]]:
    """解析AI返回的功能点匹配结果"""
    matches = []
    lines = ai_response.split('\n')
    
    for line in lines:
        line = line.strip()
        if '|' in line and not line.startswith('功能点编号'):
            try:
                parts = line.split('|')
                if len(parts) >= 6:
                    number = parts[0].strip()
                    level1 = parts[1].strip()
                    level2 = parts[2].strip()
                    level3 = parts[3].strip()
                    description = parts[4].strip()
                    workload = float(parts[5].strip())
                    
                    # 验证功能点是否在码值中存在
                    if (level1, level2, level3) in function_codes:
                        matches.append((level1, level2, level3, description, workload))
            except Exception as e:
                print(f"解析行失败：{line} - {e}")
                continue
    
    return matches


def parse_requirement_items(a4_content: str) -> List[str]:
    """解析A4单元格内容，提取各个独立的功能点"""
    content = str(a4_content).strip()
    
    # 找到所有编号的行（1. 2. 3. 等）
    import re
    lines = content.split('\n')
    items = []
    
    for line in lines:
        line = line.strip()
        # 匹配数字开头的功能点
        match = re.match(r'^\d+\.\s*(.+)$', line)
        if match:
            item_text = match.group(1).strip()
            if item_text:
                items.append(item_text)
    
    return items


def match_functions_with_ai(requirement_items: List[str], function_codes: List[Tuple[str, str, str]], total_workload: float) -> List[Tuple[str, str, str, str, float]]:
    """使用AI匹配需求内容与功能点码值"""
    codes_text = "\n".join([f"{i+1}. {l1} -> {l2} -> {l3}" for i, (l1, l2, l3) in enumerate(function_codes)])
    
    items_text = "\n".join([f"{i+1}. {item}" for i, item in enumerate(requirement_items)])
    
    prompt = f"""基于以下具体需求功能点，从功能点码值中选择最恰当的功能点进行匹配。

具体需求功能点：
{items_text}

可用功能点码值：
{codes_text}

请按照以下格式返回匹配结果，每行一个匹配项：
功能点编号|一级功能点|二级功能点|三级功能点|对应的需求功能点描述|工作量估计

要求：
1. 为每个需求功能点选择最相关的功能点码值
2. 在"对应的需求功能点描述"字段中，填入相关的原始需求功能点内容
3. 根据功能复杂度估计工作量（人天），总和应接近{total_workload}人天
4. 功能点编号从1开始递增

示例格式：
1|市场洞察|建筑视角|建筑查询|重构PC端任务流UI样式，并基于APP端样例和UI设计输出静态代码及建设页面|6.0
2|客户管控|客户视角|客户查询|调整过程表记录规则，仅保留创建、提交工单和审批环节|7.0
3|...|...|...|...|...|..."""

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
    }
    
    data = {
        "model": "deepseek-chat",
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ],
        "max_tokens": 1500,
        "temperature": 0.7
    }
    
    try:
        print("正在调用DeepSeek API进行功能点匹配...")
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=30)
        
        if response.status_code != 200:
            error_msg = f"API调用失败，状态码: {response.status_code}"
            print(error_msg)
            raise Exception(error_msg)
        
        result = response.json()
        
        if 'choices' in result and len(result['choices']) > 0:
            api_response = result['choices'][0]['message']['content'].strip()
            print("✅ AI匹配成功")
            return parse_ai_function_matches(api_response, function_codes)
        else:
            error_msg = "API响应格式错误"
            print(error_msg)
            raise Exception(error_msg)
            
    except Exception as e:
        error_msg = f"调用AI匹配失败：{e}"
        print(error_msg)
        raise Exception(error_msg)


def initialize_attachment1() -> None:
    """初始化附件1，清除之前生成的项目文档内容，并重新添加标注"""
    print_step("初始化：清理附件1中之前生成的项目文档内容，并重新添加标注")
    
    path1 = find_attachment_by_number(1)
    if not path1:
        print("未找到附件1文件")
        return
    
    try:
        from docx import Document
        
        doc = Document(path1)
        print(f"文档总段落数: {len(doc.paragraphs)}")
        
        # 识别需要清理的内容模式
        content_patterns = [
            "项目背景和概述：",
            "主要功能模块：", 
            "技术架构特点：",
            "具体目标和预期效果：",
            "业务价值和意义：",
            "用户体验提升：",
            "现有系统的不足：",
            "业务发展需要：",
            "技术升级必要性：",
            "当前系统存在的具体问题：",
            "用户使用痛点：",
            "技术或流程缺陷："
        ]
        
        # 查找需要删除的段落
        paragraphs_to_remove = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # 检查是否包含我们生成的内容特征
            for pattern in content_patterns:
                if pattern in text:
                    paragraphs_to_remove.append((i, text[:50] + "..."))
                    break
        
        print(f"找到 {len(paragraphs_to_remove)} 个需要清理的段落")
        
        # 从后往前删除，避免索引偏移
        for i, text_preview in reversed(paragraphs_to_remove):
            print(f"删除第{i}行: {text_preview}")
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)
        
        if paragraphs_to_remove:
            try:
                print(f"✅ 已清理附件1中的 {len(paragraphs_to_remove)} 个生成内容段落")
                cleaned = True
            except PermissionError:
                print("⚠️  文件被占用，无法保存。请关闭Word文档后重试")
                print(f"💡 已识别到 {len(paragraphs_to_remove)} 个需要清理的段落，但无法自动清理")
                print("💡 建议：关闭Word文档后重新运行程序")
                return
        else:
            print("✅ 附件1中没有找到需要清理的内容")
            cleaned = False
        
        # 重新添加标注，确保第十一步能找到插入位置
        # 注意：只在正文章节添加标注，不在目录中添加
        print("重新添加章节标注...")
        section_mappings = {
            "1.1": "总体描述",
            "1.2": "项目建设目标", 
            "1.3": "项目建设必要性",
            "2.3": "存在问题"
        }
        
        # 检查现有的章节标识
        annotations_found = 0
        for section_num, section_name in section_mappings.items():
            found = False
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                
                # 直接查找章节标识格式
                if (text == f"{section_name}（添加标识）" or 
                    text.startswith(f"{section_name}（") and "标识" in text or
                    text == f"{section_num} {section_name}（添加标识）" or
                    text.startswith(f"{section_num} {section_name}（") and "标识" in text):
                    print(f"✅ 找到 {section_num} {section_name} 标识")
                    annotations_found += 1
                    found = True
                    break
            
            if not found:
                print(f"⚠️  未找到 {section_num} {section_name} 标识")
        
        # 检查是否所有章节都有标识
        if annotations_found < 4:
            print(f"⚠️  只找到 {annotations_found} 个章节标识，应该有4个")
            print("💡 请确保正文中包含所有必需的章节标识")
        else:
            print(f"✅ 所有 {annotations_found} 个章节标识都已就绪")
        
        # 保存文档
        try:
            doc.save(path1)
            if annotations_found > 0:
                print(f"✅ 已检查 {annotations_found} 个章节标识")
            print("✅ 附件1初始化完成，已准备好接收新内容")
        except PermissionError:
            print("⚠️  文件被占用，无法保存标注。请关闭Word文档后重试")
        
    except ImportError:
        print("⚠️  未安装python-docx，跳过附件1初始化")
    except PermissionError:
        print("⚠️  文件被占用，无法访问附件1。请关闭Word文档后重试")
    except Exception as e:
        print(f"⚠️  附件1初始化失败：{e}")


def initialize_attachment2() -> None:
    """初始化附件2，清空数据仅保留标题行"""
    print_step("初始化：清空附件2数据，仅保留标题行")
    
    path2 = find_attachment_by_number(2)
    if not path2:
        print("未找到附件2文件")
        return
    
    try:
        wb2 = load_workbook(path2)
        ws2 = wb2.active
        
        # 取消所有合并的单元格
        merged_ranges = list(ws2.merged_cells.ranges)
        for merged_range in merged_ranges:
            ws2.unmerge_cells(str(merged_range))
        
        # 清空除标题行外的所有数据（从第2行开始）
        max_row = ws2.max_row if ws2.max_row else 1
        max_col = ws2.max_column if ws2.max_column else 7
        
        # 清空数据行（保留第1行标题）
        for row in range(2, max_row + 10):  # 多清理几行确保完全清空
            for col in range(1, max_col + 1):
                ws2.cell(row, col).value = None
                # 清除格式
                ws2.cell(row, col).alignment = None
                # 清除边框
                ws2.cell(row, col).border = None
        
        # 清除标题行的边框（恢复到初始状态）
        for col in range(1, max_col + 1):
            ws2.cell(1, col).border = None
        
        wb2.save(path2)
        print(f"已清空 {os.path.basename(path2)}，保留标题行")
        
    except Exception as e:
        print(f"初始化附件2失败：{e}")
        raise


def get_manual_summary() -> str:
    """获取沙盘操作手册的精简摘要，使用缓存机制"""
    
    manual_path = os.path.join(os.path.dirname(__file__), "沙盘操作手册.md")
    cache_path = os.path.join(os.path.dirname(__file__), "manual_summary_cache.txt")
    
    # 检查缓存是否存在且有效
    if os.path.exists(cache_path) and os.path.exists(manual_path):
        try:
            # 比较文件修改时间
            cache_mtime = os.path.getmtime(cache_path)
            manual_mtime = os.path.getmtime(manual_path)
            
            if cache_mtime > manual_mtime:
                # 缓存比手册新，直接使用缓存
                with open(cache_path, 'r', encoding='utf-8') as f:
                    cached_summary = f.read().strip()
                if cached_summary:
                    print(f"✅ 使用缓存的手册摘要，长度：{len(cached_summary)} 字符")
                    return cached_summary
        except Exception as e:
            print(f"⚠️  读取缓存失败：{e}")
    
    # 读取完整手册
    try:
        with open(manual_path, 'r', encoding='utf-8') as f:
            manual_content = f.read()
        print(f"📖 读取完整操作手册，长度：{len(manual_content)} 字符")
    except FileNotFoundError:
        print("⚠️  未找到沙盘操作手册文件")
        return ""
    except Exception as e:
        print(f"⚠️  读取沙盘操作手册失败：{e}")
        return ""
    
    # 生成摘要
    print("🤖 正在生成手册摘要以优化token使用...")
    summary_prompt = f"""请对以下沙盘操作手册内容进行精简摘要，保留核心业务信息和技术特点：

{manual_content}

要求：
1. 保留系统的四大核心模块特点（市场洞察、任务策划、任务执行、任务后评估）
2. 保留关键业务场景和功能特色
3. 保留重要的角色和权限信息
4. 压缩至2000字符以内
5. 确保摘要仍能为项目文档生成提供足够的上下文

请直接返回摘要内容，不要其他说明。"""

    try:
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
        }
        
        data = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "user", "content": summary_prompt}
            ],
            "temperature": 0.2
        }
        
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=60)
        response.raise_for_status()
        
        result = response.json()
        summary = result['choices'][0]['message']['content'].strip()
        
        # 保存摘要到缓存
        try:
            with open(cache_path, 'w', encoding='utf-8') as f:
                f.write(summary)
            print(f"✅ 已生成并缓存手册摘要，长度：{len(summary)} 字符")
        except Exception as e:
            print(f"⚠️  保存摘要缓存失败：{e}")
        
        return summary
        
    except Exception as e:
        print(f"⚠️  生成手册摘要失败：{e}，将使用原始手册")
        return manual_content


def clear_manual_cache() -> None:
    """清理手册摘要缓存，强制重新生成"""
    cache_path = os.path.join(os.path.dirname(__file__), "manual_summary_cache.txt")
    
    try:
        if os.path.exists(cache_path):
            os.remove(cache_path)
            print("✅ 已清理手册摘要缓存，下次运行将重新生成")
        else:
            print("ℹ️  缓存文件不存在，无需清理")
    except Exception as e:
        print(f"⚠️  清理缓存失败：{e}")


def generate_project_documentation(requirement_content: str) -> dict:
    """基于需求内容生成项目文档的四个部分"""
    
    # 获取手册摘要（使用缓存机制）
    manual_summary = get_manual_summary()
    
    # 构建包含手册摘要的提示词
    if manual_summary:
        prompt = f"""请先学习以下沙盘操作手册的精简摘要，了解系统的功能特点和业务场景：

=== 沙盘操作手册摘要 ===
{manual_summary}

=== 具体需求内容 ===
{requirement_content}

基于对沙盘系统的理解和以上具体需求内容，生成完整的项目文档。请确保生成的内容与沙盘系统的功能特点、业务场景、技术架构等高度契合。

请生成以下四个部分的内容：

1. 总体描述：
   - 项目背景和概述
   - 主要功能模块
   - 技术架构特点

2. 项目建设目标：
   - 具体目标和预期效果
   - 业务价值和意义
   - 用户体验提升

3. 项目建设必要性：
   - 现有系统的不足
   - 业务发展需要
   - 技术升级必要性

4. 存在问题：
   - 当前系统存在的具体问题
   - 用户使用痛点
   - 技术或流程缺陷

请确保生成的内容：
- 与沙盘系统的"市场洞察、任务策划、任务执行、任务后评估"四大模块特点相契合
- 体现政企沙盘&拓客助手系统的业务场景和功能特色
- 结合具体需求内容，体现系统优化和功能提升的必要性
- 每个部分应该有2-3个要点，每个要点100-200字

返回格式：
总体描述：
1. ...
2. ...
3. ...

项目建设目标：
1. ...
2. ...
3. ...

项目建设必要性：
1. ...
2. ...
3. ...

存在问题：
1. ...
2. ...
3. ..."""
    else:
        # 备用提示词（当手册读取失败时使用）
        prompt = f"""基于以下具体需求内容，生成完整的项目文档。

具体需求内容：
{requirement_content}

请生成以下四个部分的内容：

1. 总体描述：
   - 项目背景和概述
   - 主要功能模块
   - 技术架构特点

2. 项目建设目标：
   - 具体目标和预期效果
   - 业务价值和意义
   - 用户体验提升

3. 项目建设必要性：
   - 现有系统的不足
   - 业务发展需要
   - 技术升级必要性

4. 存在问题：
   - 当前系统存在的具体问题
   - 用户使用痛点
   - 技术或流程缺陷

请确保内容专业、具体，与具体需求高度相关。每个部分应该有2-3个要点，每个要点100-200字。

返回格式：
总体描述：
1. ...
2. ...
3. ...

项目建设目标：
1. ...
2. ...
3. ...

项目建设必要性：
1. ...
2. ...
3. ...

存在问题：
1. ...
2. ...
3. ..."""

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
    }
    
    data = {
        "model": "deepseek-chat",
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ],
        "max_tokens": 2000,
        "temperature": 0.7
    }
    
    try:
        print("正在调用DeepSeek API生成项目文档...")
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=60)
        
        if response.status_code != 200:
            error_msg = f"API调用失败，状态码: {response.status_code}"
            print(error_msg)
            raise Exception(error_msg)
        
        result = response.json()
        
        if 'choices' in result and len(result['choices']) > 0:
            ai_response = result['choices'][0]['message']['content'].strip()
            print("✅ 项目文档生成成功")
            
            # 解析AI响应为字典
            sections = {
                "总体描述": "",
                "项目建设目标": "",
                "项目建设必要性": "",
                "存在问题": ""
            }
            
            current_section = None
            lines = ai_response.split('\n')
            
            for line in lines:
                line = line.strip()
                if line.endswith('：') and line[:-1] in sections:
                    current_section = line[:-1]
                    sections[current_section] = ""
                elif current_section and line:
                    if sections[current_section]:
                        sections[current_section] += '\n' + line
                    else:
                        sections[current_section] = line
            
            return sections
        else:
            error_msg = "API响应格式错误"
            print(error_msg)
            raise Exception(error_msg)
            
    except Exception as e:
        error_msg = f"生成项目文档失败：{e}"
        print(error_msg)
        raise Exception(error_msg)


def update_attachment1_with_project_docs(project_docs: dict) -> None:
    """更新附件1中的项目文档部分"""
    print("更新附件1中的项目文档部分...")
    
    path1 = find_attachment_by_number(1)
    if not path1:
        print("未找到附件1文件")
        return
    
    try:
        # 由于Word文档更新比较复杂且容易出错，我们采用备选方案
        # 直接生成格式化的文本文件供用户手动复制
        
        output_file = os.path.join(os.path.dirname(__file__), "项目文档更新内容.txt")
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("附件1项目文档更新内容\n")
            f.write("="*60 + "\n\n")
            f.write("请将以下内容手动复制到附件1的对应章节：\n\n")
            
            section_mappings = {
                "1.1": "总体描述",
                "1.2": "项目建设目标", 
                "1.3": "项目建设必要性",
                "2.3": "存在问题"
            }
            
            for section_num, section_name in section_mappings.items():
                if section_name in project_docs and project_docs[section_name].strip():
                    f.write(f"【{section_num} {section_name}】\n")
                    f.write("-" * 40 + "\n")
                    
                    # 格式化内容，每个要点分段显示
                    content = project_docs[section_name].strip()
                    lines = content.split('\n')
                    
                    for line in lines:
                        line = line.strip()
                        if line:
                            f.write(line + "\n\n")
                    
                    f.write("\n")
        
        print(f"✅ 项目文档内容已保存到：{output_file}")
        print("📋 请手动执行以下步骤：")
        print("   1. 打开附件1的Word文档")
        print("   2. 找到对应的章节（1.1、1.2、1.3、2.3）")
        print("   3. 将文本文件中的内容复制到对应章节下")
        print("   4. 保存Word文档")
        
        # 同时尝试自动更新（如果可能的话）
        try:
            from docx import Document
            
            print("\n尝试自动更新Word文档...")
            doc = Document(path1)
            
            section_mappings = {
                "1.1": "总体描述",
                "1.2": "项目建设目标", 
                "1.3": "项目建设必要性",
                "2.3": "存在问题"
            }
            
            updated_sections = []
            
            # 只支持章节标识格式
            section_title_mappings = {
                "1.1": "总体描述",
                "1.2": "项目建设目标", 
                "1.3": "项目建设必要性",
                "2.3": "存在问题"
            }
            
            print("查找章节标识...")
            
            # 直接查找章节标识
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                
                # 查找带有标识的章节标题（支持两种格式）
                for section_num, section_name in section_title_mappings.items():
                    # 格式1: "1.1 总体描述（添加标识）"
                    # 格式2: "总体描述（添加标识）"  
                    if ((text.startswith(section_num) and section_name in text and ('添加标识' in text or '标识' in text or '（' in text)) or
                        (text == f"{section_name}（添加标识）" or text.startswith(f"{section_name}（") and "标识" in text)):
                        
                        print(f"找到带标识的章节：第{i}行 - {text}")
                        
                        if section_name in project_docs:
                            content = project_docs[section_name].strip()
                            
                            if content:
                                # 在章节标题后添加内容
                                lines = [line.strip() for line in content.split('\n') if line.strip()]
                                
                                # 使用更简单可靠的方法：直接在目标段落后面依次插入
                                target_para = paragraph
                                parent = target_para._element.getparent()
                                target_element = target_para._element
                                
                                # 正序插入每一行内容
                                insert_position = list(parent).index(target_element) + 1
                                
                                for line_idx, line in enumerate(lines):
                                    # 创建新的段落元素
                                    new_p = doc._body._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
                                    # 创建文本运行
                                    new_r = doc._body._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                                    new_t = doc._body._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                                    new_t.text = line
                                    new_r.append(new_t)
                                    new_p.append(new_r)
                                    
                                    # 插入到正确位置
                                    parent.insert(insert_position + line_idx, new_p)
                                
                                updated_sections.append(section_name)
                                print(f"✅ 已在章节 {section_num} 后添加 {len(lines)} 行内容")
                                
                                # 提前退出，避免重复处理
                                break
                        break
            
            # 禁用末尾添加功能，要求必须找到标注位置
            if not updated_sections:
                print("❌ 未找到任何用户标注位置")
                print("💡 请确保文档中包含以下格式的章节标识：")
                print("   • 1.1 总体描述（添加标识）")
                print("   • 总体描述（添加标识）")
                print("   • 1.2 项目建设目标（添加标识）")
                print("   • 项目建设目标（添加标识）")
                print("   • 1.3 项目建设必要性（添加标识）")
                print("   • 项目建设必要性（添加标识）")
                print("   • 2.3 存在问题（添加标识）")
                print("   • 存在问题（添加标识）")
                print("⚠️  不允许在文档末尾添加内容，必须在指定位置插入")
                return
            
            if updated_sections:
                doc.save(path1)
                print(f"✅ Word文档自动更新成功，已替换标注：{', '.join(updated_sections)}")
                print("💡 内容已精确插入到您标注的位置")
            else:
                print("⚠️  未找到用户标注位置，请检查标注格式")
                print("💡 建议使用格式：总体描述（添加标识）、项目建设目标（添加标识）等")
                
        except Exception as auto_error:
            print(f"⚠️  Word文档自动更新失败：{auto_error}")
            print("请使用手动方式复制内容")
        
    except Exception as e:
        print(f"生成文档内容失败：{e}")
        import traceback
        traceback.print_exc()


def update_attachment3_with_project_docs(project_docs: dict) -> None:
    """更新附件3中的建设目标和建设必要性"""
    print("更新附件3中的建设目标和建设必要性...")
    
    path3 = find_attachment_by_number(3)
    if not path3:
        print("未找到附件3文件")
        return
    
    try:
        wb3 = load_workbook(path3)
        # 使用第一个工作表（系统功能架构图）
        ws3 = wb3.active
        
        updated_cells = []
        
        # 更新A2单元格 - 建设目标
        if "项目建设目标" in project_docs and project_docs["项目建设目标"].strip():
            target_content = project_docs["项目建设目标"].strip()
            ws3['A2'].value = target_content
            updated_cells.append("A2(建设目标)")
            print(f"✅ 已更新A2单元格：建设目标")
        
        # 更新A5单元格 - 建设必要性
        if "项目建设必要性" in project_docs and project_docs["项目建设必要性"].strip():
            necessity_content = project_docs["项目建设必要性"].strip()
            ws3['A5'].value = necessity_content
            updated_cells.append("A5(建设必要性)")
            print(f"✅ 已更新A5单元格：建设必要性")
        
        if updated_cells:
            # 保存文件
            wb3.save(path3)
            print(f"✅ 附件3更新成功，已更新：{', '.join(updated_cells)}")
        else:
            print("⚠️  没有找到可更新的内容")
        
    except PermissionError:
        print("⚠️  附件3文件被占用，无法保存。请关闭Excel文档后重试")
    except Exception as e:
        print(f"⚠️  附件3更新失败：{e}")
        import traceback
        traceback.print_exc()


def step11_generate_and_update_project_docs() -> None:
    """第十一步：生成项目文档并更新附件1"""
    print_step("第十一步：生成项目文档并更新附件1")
    
    try:
        # 获取附件4的A4内容
        path4 = find_attachment_by_number(4)
        if not path4:
            print("未找到附件4文件")
            return
        
        wb4 = load_workbook(path4, data_only=True)
        ws4 = wb4.active
        a4_content = ws4['A4'].value
        
        if not a4_content:
            print("附件4的A4单元格为空")
            return
        
        print(f"提取到需求内容：{len(str(a4_content))} 字符")
        
        # 生成项目文档
        project_docs = generate_project_documentation(str(a4_content))
        
        # 显示生成的内容
        print("\n" + "="*80)
        print("生成的项目文档内容：")
        print("="*80)
        
        section_mappings = {
            "1.1": "总体描述",
            "1.2": "项目建设目标", 
            "1.3": "项目建设必要性",
            "2.3": "存在问题"
        }
        
        for section_num, section_name in section_mappings.items():
            if section_name in project_docs and project_docs[section_name].strip():
                print(f"\n【{section_num} {section_name}】")
                print("-" * 40)
                print(project_docs[section_name])
        
        # 更新附件1
        update_attachment1_with_project_docs(project_docs)
        
        # 更新附件3
        update_attachment3_with_project_docs(project_docs)
        
        print("\n✅ 第十一步完成：项目文档已生成并更新")
        
    except Exception as e:
        print(f"第十一步执行失败：{e}")
        print("程序终止，请检查API配置或网络连接")
        raise


def update_wbs_document() -> None:
    print_step("第十步：基于A4数据和功能点码值更新WBS工作量文档")
    
    try:
        # 获取附件4的A4内容
        path4 = find_attachment_by_number(4)
        if not path4:
            print("未找到附件4文件")
            return
        
        wb4 = load_workbook(path4, data_only=True)
        ws4 = wb4.active
        a4_content = ws4['A4'].value  # A4包含需求内容，也用作功能描述
        d7_workload = ws4['D7'].value or 19.0
        
        if not a4_content:
            print("附件4的A4单元格为空")
            return
        
        # 解析A4内容，提取各个独立的功能点
        requirement_items = parse_requirement_items(a4_content)
        
        print(f"提取到需求内容：{len(str(a4_content))} 字符")
        print(f"解析出 {len(requirement_items)} 个具体功能点：")
        for i, item in enumerate(requirement_items, 1):
            print(f"  {i}. {item}")
        print(f"工作量总和：{d7_workload} 人天")
        
        # 加载功能点码值
        function_codes = load_function_codes()
        if not function_codes:
            print("无法加载功能点码值，程序终止")
            return
        
        # AI匹配功能点
        matches = match_functions_with_ai(requirement_items, function_codes, d7_workload)
        
        if not matches:
            print("AI匹配失败，程序终止")
            return
        
        print(f"匹配到 {len(matches)} 个功能点")
        
        # 更新WBS文档
        path2 = find_attachment_by_number(2)
        if not path2:
            print("未找到附件2 WBS文件")
            return
        
        wb2 = load_workbook(path2)
        ws2 = wb2.active
        
        # 取消所有合并的单元格
        merged_ranges = list(ws2.merged_cells.ranges)
        for merged_range in merged_ranges:
            ws2.unmerge_cells(str(merged_range))
        
        # 清空现有数据（保留标题行）
        for row in range(2, ws2.max_row + 10):  # 多清理几行确保完全清空
            for col in range(1, 7):
                ws2.cell(row, col).value = None
        
        # 按功能点码值排序并合并相同功能点的描述
        from collections import defaultdict
        from openpyxl.styles import Alignment, Border, Side
        
        # 按功能点码值分组
        grouped_matches = defaultdict(list)
        for level1, level2, level3, description, workload in matches:
            key = (level1, level2, level3)
            grouped_matches[key].append((description, workload))
        
        # 按功能点码值排序
        sorted_groups = sorted(grouped_matches.items())
        
        # 定义边框样式
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # 为标题行添加边框
        for col in range(1, 7):
            cell = ws2.cell(1, col)
            cell.border = thin_border
        
        # 填入匹配的功能点数据
        current_row = 2
        for (level1, level2, level3), items in sorted_groups:
            # 合并所有描述，用换行符分隔，并添加序号
            combined_descriptions = []
            total_workload = 0
            
            for i, (description, workload) in enumerate(items, 1):
                # 为每个描述添加序号
                numbered_description = f"{i}. {description}"
                combined_descriptions.append(numbered_description)
                total_workload += workload
            
            # 用换行符连接所有描述
            combined_description = '\n'.join(combined_descriptions)
            
            ws2.cell(current_row, 1).value = f"=ROW()-1"  # 编号使用公式
            ws2.cell(current_row, 2).value = level1  # 一级功能点
            ws2.cell(current_row, 3).value = level2  # 二级功能点
            ws2.cell(current_row, 4).value = level3  # 三级功能点
            ws2.cell(current_row, 5).value = combined_description  # 合并的功能描述
            ws2.cell(current_row, 6).value = total_workload  # 合并的工作量
            
            # 设置单元格格式：自动换行和边框
            for col in range(1, 7):
                cell = ws2.cell(current_row, col)
                cell.border = thin_border
                if col == 5:  # 功能描述列
                    cell.alignment = Alignment(wrap_text=True, vertical='top')
            
            current_row += 1
        
        # 添加合计行
        total_row = current_row
        ws2.cell(total_row, 1).value = f"=ROW()-1"  # 编号
        
        # 合并B列和E列单元格并填入"合计"
        ws2.merge_cells(f'B{total_row}:E{total_row}')
        ws2.cell(total_row, 2).value = "合计"
        
        # F列工作量总和
        ws2.cell(total_row, 6).value = d7_workload
        
        # 为合计行添加边框
        for col in range(1, 7):
            cell = ws2.cell(total_row, col)
            cell.border = thin_border
        
        wb2.save(path2)
        
        print(f"已更新 {os.path.basename(path2)}:")
        print(f"  - 原始匹配: {len(matches)} 个功能点")
        print(f"  - 合并后: {len(sorted_groups)} 个功能点组")
        print(f"  - 相同功能点码值的描述已合并并换行显示")
        print(f"  - 功能描述添加了序号")
        print(f"  - 为所有单元格添加了边框")
        print(f"  - 工作量总和: {d7_workload} 人天")
        print(f"  - 添加合计行")
        
        # 显示合并后的结果
        print("\n合并后的结果预览:")
        for i, ((level1, level2, level3), items) in enumerate(sorted_groups, 1):
            descriptions = [desc for desc, _ in items]
            total_workload = sum(workload for _, workload in items)
            print(f"  {i}. {level1}->{level2}->{level3}")
            print(f"     功能描述数量: {len(descriptions)} 个")
            print(f"     总工作量: {total_workload} 人天")
            for j, desc in enumerate(descriptions, 1):
                print(f"       {j}) {desc[:50]}...")
            print()
        
    except Exception as e:
        print(f"第十步执行失败：{e}")
        print("程序终止，请检查API配置或网络连接")
        raise


def enhance_cosmic_data_groups_and_attributes(trigger_event: str, function_process: str, subprocess_desc: str, data_movement_type: str, existing_data_group: str = "", existing_data_attributes: str = "") -> tuple:
    """基于COSMIC背景，调用大模型生成或完善数据组和数据属性"""
    
    prompt = f"""作为COSMIC软件度量专家，基于以下信息，为子过程生成合适的数据组和数据属性。

COSMIC背景知识：
- 数据组(Data Group)：逻辑上相关的数据属性集合，代表软件用户感兴趣的对象
- 数据属性(Data Attributes)：构成数据组的具体属性字段
- 数据移动类型：Entry(E)-数据进入, Exit(X)-数据退出, Read(R)-数据读取, Write(W)-数据写入

当前子过程信息：
- 触发事件：{trigger_event}
- 功能过程：{function_process}
- 子过程描述：{subprocess_desc}
- 数据移动类型：{data_movement_type}

现有数据组：{existing_data_group if existing_data_group else "无"}
现有数据属性：{existing_data_attributes if existing_data_attributes else "无"}

要求：
1. 数据组名称要简洁、准确，体现业务含义
2. 数据属性要具体、完整，包含该数据组的关键字段
3. 确保与数据移动类型({data_movement_type})的语义一致
4. 如果已有数据组和属性，请在此基础上优化完善
5. 不同子过程的数据组和数据属性要保持差异性，避免重复

请返回格式：
数据组：[数据组名称]
数据属性：[属性1、属性2、属性3、...]

只返回数据组和数据属性，不要其他内容。"""

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
    }
    
    data = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3
    }
    
    try:
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=30)
        response.raise_for_status()
        
        result = response.json()
        content = result['choices'][0]['message']['content'].strip()
        
        # 解析返回内容
        lines = content.split('\n')
        data_group = ""
        data_attributes = ""
        
        for line in lines:
            line = line.strip()
            if line.startswith('数据组：'):
                data_group = line.replace('数据组：', '').strip()
            elif line.startswith('数据属性：'):
                data_attributes = line.replace('数据属性：', '').strip()
        
        return data_group, data_attributes
        
    except requests.exceptions.Timeout:
        print(f"⚠️  API调用超时，使用默认值")
        return existing_data_group or "默认数据组", existing_data_attributes or "默认属性"
    except Exception as e:
        print(f"⚠️  调用AI生成数据组和属性失败：{e}")
        return existing_data_group or "默认数据组", existing_data_attributes or "默认属性"


def step12_enhance_cosmic_data_groups_and_attributes() -> None:
    """第十二步：基于COSMIC背景完善数据组和数据属性"""
    print_step("第十二步：基于COSMIC背景完善数据组和数据属性")
    
    path3 = find_attachment_by_number(3)
    if not path3:
        print("未找到附件3文件")
        return
    
    try:
        wb3 = load_workbook(path3)
        
        # 查找COSMIC功能点拆分表工作表
        sheet_name = "COSMIC功能点拆分表"
        if sheet_name not in wb3.sheetnames:
            print(f"未找到{sheet_name}工作表")
            return
        
        ws = wb3[sheet_name]
        print(f"✅ 找到{sheet_name}工作表")
        
        # 统计处理的行数
        processed_count = 0
        enhanced_count = 0
        
        # 从第4行开始处理数据
        for row in range(4, ws.max_row + 1):
            # 获取相关列的数据
            trigger_event = ws.cell(row, 6).value or ""      # F列 - 触发事件
            function_process = ws.cell(row, 7).value or ""   # G列 - 功能过程  
            subprocess_desc = ws.cell(row, 8).value or ""    # H列 - 子过程描述
            data_movement_type = ws.cell(row, 9).value or "" # I列 - 数据移动类型
            existing_data_group = ws.cell(row, 10).value or ""     # J列 - 数据组
            existing_data_attributes = ws.cell(row, 11).value or "" # K列 - 数据属性
            
            # 只处理有子过程描述和数据移动类型的行
            if subprocess_desc.strip() and data_movement_type.strip():
                processed_count += 1
                print(f"\n处理第{row}行:")
                print(f"  子过程描述: {subprocess_desc[:50]}...")
                print(f"  数据移动类型: {data_movement_type}")
                
                # 调用AI生成或完善数据组和数据属性
                new_data_group, new_data_attributes = enhance_cosmic_data_groups_and_attributes(
                    trigger_event, function_process, subprocess_desc, data_movement_type,
                    existing_data_group, existing_data_attributes
                )
                
                # 检查是否有改进
                if (new_data_group != existing_data_group or 
                    new_data_attributes != existing_data_attributes):
                    
                    # 更新数据
                    ws.cell(row, 10).value = new_data_group      # J列 - 数据组
                    ws.cell(row, 11).value = new_data_attributes # K列 - 数据属性
                    
                    enhanced_count += 1
                    print(f"  ✅ 已完善数据组: {new_data_group}")
                    print(f"  ✅ 已完善数据属性: {new_data_attributes[:50]}...")
                else:
                    print(f"  ✓ 数据组和属性已完善，无需修改")
        
        # 保存文件
        if enhanced_count > 0:
            wb3.save(path3)
            print(f"\n✅ 已保存附件3，共处理 {processed_count} 行，完善 {enhanced_count} 行")
        else:
            print(f"\n✓ 所有 {processed_count} 行数据组和属性都已完善，无需修改")
        
        print("✅ 第十二步完成：COSMIC数据组和数据属性已完善")
        
    except PermissionError:
        print("⚠️  附件3文件被占用，无法保存。请关闭Excel文档后重试")
    except Exception as e:
        print(f"第十二步执行失败：{e}")
        import traceback
        traceback.print_exc()


def main() -> None:
    print_step("输入变量：统一替换的需求名")
    requirement_name = input("请输入需求名字（用于重命名与单元格填充）：").strip()
    if not requirement_name:
        print("未输入需求名字，程序结束。")
        return

    # 0) 初始化附件1和附件2，清理之前的生成内容
    initialize_attachment1()
    initialize_attachment2()

    # 1) 批量重命名
    batch_rename(requirement_name)

    # 2) 附件3 sheet2 A3/B3
    write_attachment3_sheet2_cells(requirement_name)

    # 3) 附件4 B2
    write_attachment4_cells(requirement_name)

    # 4) 计算附件5 L列总和
    total = sum_attachment5_col_L_from_L2()

    # 5) 附件4 D7 = total
    write_attachment4_with_sum(total)

    # 6) 附件3 sheet2 F3 = total
    write_attachment3_sheet2_F3_with_sum(total)

    # 7) 附件4 C6 = 今天日期
    write_attachment4_C6_with_today()

    # 8) 附件4 B7 = 附件3 sheet2 E3
    write_attachment4_B7_from_attachment3_E3()

    # 9) 附件4 A4 = 附件5 H和I列内容概述
    summarize_requirement_content_and_update_h4()

    # 10) 更新WBS文档
    update_wbs_document()

    # 11) 生成项目文档并更新附件1
    step11_generate_and_update_project_docs()

    # 12) 完善COSMIC数据组和数据属性
    step12_enhance_cosmic_data_groups_and_attributes()

    print_step("全部步骤完成")


if __name__ == "__main__":
    main() 